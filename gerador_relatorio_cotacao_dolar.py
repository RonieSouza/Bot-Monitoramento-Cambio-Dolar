from selenium.webdriver import Chrome
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import NoSuchElementException, ElementNotVisibleException, ElementNotSelectableException
from selenium.webdriver.support import expected_conditions as EC
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from docx2pdf import convert
import pytz
import os

class Config:
    """Configurações do projeto."""

    CHROME_OPTIONS = {
        'window_size': '--window-size=1360,750',
        'incognito': '--incognito',
        'headless': '--headless',
        'disable_gpu': '--disable-gpu',
        'notifications': {
            "profile.default_content_setting_values.notifications": 2
        }
    }
    XPATH_PRECO = '//div[@class="YMlKec fxKbKc"]'
    XPATH_DATA = '//div[@class="ygUjEc"]'
    CAMINHO_CAPTURA_TELA = "screenshot.png"
    CAMINHO_DOC = "cotacao_dolar.docx"
    CAMINHO_PDF = os.path.join(os.getcwd(),"cotacao_dolar.pdf")
    URL = 'https://www.google.com/finance/quote/USD-BRL'

class DriverConfig:
    """Configura o driver do navegador."""
    
    @staticmethod
    def configurar_driver():
        """Configura e retorna o driver do Chrome e o WebDriverWait."""
        chrome_options = Options()
        DriverConfig._adicionar_opcoes(chrome_options)
        driver = Chrome(options=chrome_options)
        wait = WebDriverWait(driver, 10, poll_frequency=1,
                            ignored_exceptions=[NoSuchElementException, ElementNotVisibleException, ElementNotSelectableException])
        return driver, wait

    @staticmethod
    def _adicionar_opcoes(chrome_options):
        """Adiciona as opções ao driver do Chrome."""
        for key, option in Config.CHROME_OPTIONS.items():
            if isinstance(option, dict):
                chrome_options.add_experimental_option("prefs", option)
            else:
                chrome_options.add_argument(option)

class DateFormatter:
    """Responsável pela formatação de datas."""

    @staticmethod
    def formatar_data(data_em_texto):
        """Formata a data extraída do site para o formato desejado."""
        try:
            date_str = data_em_texto.split(' UTC')[0]
            date_str_with_year = f"{date_str}, {datetime.now().year}"
            original_date = datetime.strptime(date_str_with_year, "%b %d, %I:%M:%S %p, %Y")
            local_date = original_date.replace(tzinfo=pytz.UTC).astimezone(pytz.timezone('America/Sao_Paulo'))
            return local_date.strftime("%d/%m/%Y %H:%M:%S")
        except ValueError as e:
            raise ValueError(f"Erro ao formatar a data: {e}")

class ReportGenerator:
    """Gera o relatório em formatos Word e PDF."""

    def __init__(self, url, preco, data_hora_formatada):
        self.url = url
        self.preco = preco
        self.data_hora_formatada = data_hora_formatada

    def criar_documento(self):
        """Cria o documento Word e o converte para PDF."""
        try:
            doc = Document()
            self._adicionar_cabecalho(doc)
            self._adicionar_informacoes(doc)
            self._adicionar_imagem(doc, Config.CAMINHO_CAPTURA_TELA)
            doc.save(Config.CAMINHO_DOC)
            convert(Config.CAMINHO_DOC, Config.CAMINHO_PDF)
        except Exception as e:
            raise RuntimeError(f"Erro ao criar o documento: {e}")

    def _adicionar_cabecalho(self, doc):
        """Adiciona o cabeçalho ao documento."""
        doc.add_paragraph('Relatório de Cotação', style='Title')
        doc.add_heading('Cotação do Dólar', level=1)

    def _adicionar_informacoes(self, doc):
        """Adiciona as informações principais ao documento."""
        paragraph = doc.add_paragraph()
        paragraph.add_run('O dólar está no valor de ').bold = False
        paragraph.add_run(self.preco).bold = True
        paragraph.add_run(', na data ').bold = False
        paragraph.add_run(self.data_hora_formatada).bold = True
        paragraph.add_run('\nValor cotado no site: ').bold = False
        url_run = paragraph.add_run(self.url)
        url_run.font.color.rgb = RGBColor(0, 0, 255)
        url_run.font.underline = True
        url_run.font.italic = True

    def _adicionar_imagem(self, doc, caminho_imagem):
        """Adiciona a captura de tela ao documento."""
        doc.add_picture(caminho_imagem, width=Inches(6))
        legenda = doc.add_paragraph('Captura da tela mostrando a cotação atual do dólar.')
        legenda.alignment = 1

def main():
    """Função principal que executa o fluxo do programa."""

    print("\033[33m\033[1mIniciando o processo de coleta de dados...\033[0m")

    driver = None
    erro_ocorrido = False

    try:
        driver, wait = DriverConfig.configurar_driver()
        driver.get(Config.URL)

        preco = wait.until(EC.visibility_of_element_located((By.XPATH, Config.XPATH_PRECO))).text
        data_hora = wait.until(EC.visibility_of_element_located((By.XPATH, Config.XPATH_DATA))).text

        data_hora_formatada = DateFormatter.formatar_data(data_hora)
        driver.save_screenshot(Config.CAMINHO_CAPTURA_TELA)

        report_generator = ReportGenerator(
            url=Config.URL,
            preco=preco,
            data_hora_formatada=data_hora_formatada
        )
        report_generator.criar_documento()

    except RuntimeError as e:
        if "No such file or directory" in str(e):
            print("\033[31m\033[1mErro ao criar o documento. Por favor, execute o programa como administrador, pois o diretório onde o programa está instalado pode estar protegido.\033[0m")
        else:
            print(f"Erro durante a execução do programa: {e}")
        erro_ocorrido = True

    except Exception as e:
        print(f"Erro durante a execução do programa: {e}")
        erro_ocorrido = True

    finally:
        if driver:
            driver.quit()

    if not erro_ocorrido:
        print(f"\033[32m\033[1mProcesso finalizado com sucesso. Arquivo -> {Config.CAMINHO_PDF}\033[0m")

    input('Aperte a tecla Enter para fechar')

if __name__ == "__main__":
    main()